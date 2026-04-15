"""Analyze the reference template's styling in detail."""
from docx import Document
from docx.shared import Pt, Emu

doc = Document('../KPMG_Vinod Kumar_DBT & Data Engineer_9+ years_Bengaluru (1).docx')

# Page margins
s = doc.sections[0]
print("=== PAGE SETUP ===")
print(f"  Top margin: {s.top_margin} ({Emu(s.top_margin).inches:.2f} in)")
print(f"  Bottom margin: {s.bottom_margin} ({Emu(s.bottom_margin).inches:.2f} in)")
print(f"  Left margin: {s.left_margin} ({Emu(s.left_margin).inches:.2f} in)")
print(f"  Right margin: {s.right_margin} ({Emu(s.right_margin).inches:.2f} in)")
print(f"  Page width: {s.page_width}")
print(f"  Page height: {s.page_height}")
print(f"  Header distance: {s.header_distance}")
print(f"  Footer distance: {s.footer_distance}")

# Header/footer
print("\n=== HEADER ===")
for h in s.header.paragraphs:
    print(f"  Header para: '{h.text[:60]}' style={h.style.name}")
    for r in h.runs:
        print(f"    Run: '{r.text[:40]}' font={r.font.name} size={r.font.size} bold={r.font.bold}")

print("\n=== FOOTER ===")
for f_para in s.footer.paragraphs:
    print(f"  Footer para: '{f_para.text[:60]}' style={f_para.style.name}")

# Paragraph styles
print("\n=== PARAGRAPHS (first 25) ===")
for i, para in enumerate(doc.paragraphs[:25]):
    pf = para.paragraph_format
    print(f"\n  [{i}] Style: {para.style.name}")
    print(f"       Text: '{para.text[:80]}'")
    print(f"       Alignment: {pf.alignment}")
    print(f"       Space before: {pf.space_before}")
    print(f"       Space after: {pf.space_after}")
    print(f"       Line spacing: {pf.line_spacing}")
    print(f"       Left indent: {pf.left_indent}")
    for j, run in enumerate(para.runs[:3]):
        f = run.font
        print(f"       Run[{j}]: font={f.name} size={f.size} bold={f.bold} italic={f.italic} color={f.color.rgb if f.color and f.color.rgb else None} underline={f.underline}")

# Check for numbering / bullet formats
print("\n=== NUMBERING INFO ===")
from lxml import etree
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for i, para in enumerate(doc.paragraphs[:25]):
    numPr = para._element.find('.//w:numPr', ns)
    if numPr is not None:
        ilvl = numPr.find('w:ilvl', ns)
        numId = numPr.find('w:numId', ns)
        level = ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if ilvl is not None else '?'
        nid = numId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if numId is not None else '?'
        print(f"  [{i}] numId={nid} ilvl={level} text='{para.text[:50]}'")
