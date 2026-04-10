from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import copy


class TemplateEngine:
    """Engine for applying template styling and structure to generated documents"""

    def __init__(self, template_path: str):
        """
        Initialize template engine with reference document.
        
        Args:
            template_path: Path to reference template DOCX
        """
        self.template_doc = Document(template_path)
        self.template_styles = self._extract_styles()
        self.template_formatting = self._extract_formatting()

    def _safe_set_style(self, para, style_name: str) -> None:
        """Apply a style only when it exists in the document."""
        try:
            if style_name in self.template_doc.styles:
                para.style = style_name
        except Exception:
            pass

    def _extract_styles(self) -> Dict[str, Any]:
        """Extract all styles from template document"""
        styles = {}
        try:
            for style in self.template_doc.styles:
                styles[style.name] = {
                    "font_name": style.font.name or "Calibri",
                    "font_size": style.font.size.pt if style.font.size else 11,
                    "bold": style.font.bold or False,
                    "italic": style.font.italic or False,
                    "underline": style.font.underline or False,
                    "color": style.font.color.rgb if style.font.color.rgb else None,
                }
        except:
            pass
        return styles

    def _extract_formatting(self) -> Dict[str, Any]:
        """Extract document formatting from template"""
        formatting = {
            "margins": {},
            "sections": [],
            "numbering": {}
        }

        try:
            # Extract margins
            section = self.template_doc.sections[0]
            formatting["margins"] = {
                "top": section.top_margin,
                "bottom": section.bottom_margin,
                "left": section.left_margin,
                "right": section.right_margin
            }

            # Extract section structure
            for para in self.template_doc.paragraphs:
                formatting["sections"].append({
                    "style": para.style.name,
                    "level": para.paragraph_format.left_indent,
                    "space_before": para.paragraph_format.space_before,
                    "space_after": para.paragraph_format.space_after,
                })
        except:
            pass

        return formatting

    def apply_template_to_doc(self, target_doc: Document) -> Document:
        """
        Apply template styles and formatting to target document.
        
        Args:
            target_doc: Document to apply template to
            
        Returns:
            Document with template applied
        """
        try:
            # Copy margins from template
            target_section = target_doc.sections[0]
            source_section = self.template_doc.sections[0]

            target_section.top_margin = source_section.top_margin
            target_section.bottom_margin = source_section.bottom_margin
            target_section.left_margin = source_section.left_margin
            target_section.right_margin = source_section.right_margin

            # Apply heading styles
            self._apply_heading_styles(target_doc)

            # Apply numbering styles
            self._apply_numbering(target_doc)

        except Exception as e:
            raise Exception(f"Error applying template: {e}")

        return target_doc

    def _apply_heading_styles(self, doc: Document):
        """Apply heading styles from template"""
        try:
            for para in doc.paragraphs:
                current_style = getattr(para.style, "name", "")
                self._safe_set_style(para, current_style)

        except Exception as e:
            print(f"Error applying heading styles: {e}")

    def _apply_numbering(self, doc: Document):
        """Apply proper numbering from template"""
        try:
            for para in doc.paragraphs:
                style_name = getattr(para.style, "name", "")
                if style_name:
                    self._safe_set_style(para, style_name)

        except Exception as e:
            print(f"Error applying numbering: {e}")

    def copy_template_structure(self) -> Document:
        """
        Create a new document as copy of template.
        
        Returns:
            New document with template structure
        """
        return copy.deepcopy(self.template_doc)

    def get_template_sections(self) -> list:
        """Get section headers from template"""
        sections = []
        try:
            for para in self.template_doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    sections.append(para.text)
        except:
            pass
        return sections
