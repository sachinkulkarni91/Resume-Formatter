from docx import Document
from lxml import etree

doc = Document('../KPMG_Vinod Kumar_DBT & Data Engineer_9+ years_Bengaluru (1).docx')
s = doc.sections[0]

# Check header for images
header_xml = etree.tostring(s.header._element, pretty_print=True).decode()
has_image = 'blip' in header_xml.lower() or 'image' in header_xml.lower()
print(f'Header has image: {has_image}')
print(f'Header paragraphs: {len(s.header.paragraphs)}')
for p in s.header.paragraphs:
    print(f'  Header text: {repr(p.text[:40])}')

# Shading on paragraphs (gray background bars)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for i, para in enumerate(doc.paragraphs[:10]):
    shd = para._element.find('.//w:shd', ns)
    if shd is not None:
        fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        print(f'Shading on para [{i}] fill={fill} text={para.text[:50]}')
    # Also check pPr > shd
    pPr = para._element.find('w:pPr', ns)
    if pPr is not None:
        shd2 = pPr.find('w:shd', ns)
        if shd2 is not None:
            fill2 = shd2.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
            print(f'pPr Shading on para [{i}] fill={fill2} text={para.text[:50]}')
