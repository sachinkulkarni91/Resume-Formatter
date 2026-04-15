from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
import copy
import os


class DocxGenerator:
    """Generator for creating formatted DOCX documents that closely
    reproduce the styling of a given KPMG-style reference template."""

    # -- Formatting constants extracted from the reference template --
    FONT_NAME = "Calibri"
    BODY_SIZE = Pt(10.5)           
    NAME_SIZE = Pt(14)           
    TITLE_SIZE = Pt(10.5)
    HEADING_SIZE = Pt(11)
    JOB_TITLE_SIZE = Pt(10.5)
    SHADING_COLOR = "D9D9D9"     # Gray background on section headings

    def __init__(self, template_path: str = None):
        """
        Initialize document generator.

        Args:
            template_path: Optional path to template document
        """
        self.template_path = template_path

        template_ext = os.path.splitext(template_path)[1].lower() if template_path else ""

        if template_path and template_ext in [".docx", ".doc"]:
            self.template_doc = Document(template_path)
            # Try to detect the font used in the template
            self._detect_template_font()
            # Extract page margins & size
            self._template_margins = self._extract_margins()
        else:
            self.template_doc = None
            self._template_margins = None

        # Working document â€“ always fresh
        self.doc = Document()

    # ------------------------------------------------------------------
    # Template analysis helpers
    # ------------------------------------------------------------------

    def _detect_template_font(self):
        """Detect the primary font family used in the template."""
        for para in self.template_doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    self.FONT_NAME = run.font.name
                    return

    def _extract_margins(self) -> dict:
        """Read page margins from the template."""
        try:
            s = self.template_doc.sections[0]
            return {
                "top": s.top_margin,
                "bottom": s.bottom_margin,
                "left": s.left_margin,
                "right": s.right_margin,
                "width": s.page_width,
                "height": s.page_height,
                "header_distance": s.header_distance,
                "footer_distance": s.footer_distance,
            }
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Low-level formatting helpers
    # ------------------------------------------------------------------

    def _set_run_font(self, run, size=None, bold=None, italic=None, color=None):
        """Apply font formatting to a run."""
        run.font.name = self.FONT_NAME
        # Also set the East-Asian / complex-script font so Word respects it
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
        if size:
            run.font.size = size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def _set_paragraph_spacing(self, para, space_before=None, space_after=None,
                                line_spacing=None):
        """Set paragraph spacing."""
        pf = para.paragraph_format
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after
        if line_spacing is not None:
            pf.line_spacing = line_spacing

    def _add_shading(self, para, color: str):
        """Add gray background shading to a paragraph (like the KPMG section bars)."""
        pPr = para._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        pPr.append(shd)

    def _add_bullet_paragraph(self, text: str) -> "Paragraph":
        """Add a paragraph formatted as a bullet point."""
        para = self.doc.add_paragraph(style="List Bullet")
        para.clear()
        run = para.add_run(str(text))
        self._set_run_font(run, size=self.BODY_SIZE)
        self._set_paragraph_spacing(para, space_after=Emu(8890), line_spacing=1.03)
        return para

    # ------------------------------------------------------------------
    # Document generation
    # ------------------------------------------------------------------

    def generate_from_structured_data(self, data: Dict[str, Any]) -> Document:
        """
        Generate document from structured resume data.

        Args:
            data: Structured resume data

        Returns:
            Generated Document
        """
        # Always create a fresh document
        self.doc = Document()

        # Apply template page layout
        if self._template_margins:
            s = self.doc.sections[0]
            m = self._template_margins
            s.top_margin = m["top"]
            s.bottom_margin = m["bottom"]
            s.left_margin = m["left"]
            s.right_margin = m["right"]
            s.page_width = m["width"]
            s.page_height = m["height"]
            s.header_distance = m["header_distance"]
            s.footer_distance = m["footer_distance"]

        # Remove the default empty paragraph that Document() creates
        if self.doc.paragraphs:
            p = self.doc.paragraphs[0]._element
            p.getparent().remove(p)

        self._add_header_logo()

        # Build resume sections
        if "name" in data:
            self._add_name_section(data["name"], data.get("title", ""))

        if "summary" in data and data["summary"]:
            self._add_heading("PROFESSIONAL SUMMARY")
            summary_data = data["summary"]

            # Handle whatever format the LLM gives us, keeping it as a paragraph
            if isinstance(summary_data, str):
                self._add_body_text(summary_data)
            elif isinstance(summary_data, dict):
                bg = summary_data.get("background", "").strip()
                bullets = summary_data.get("bullets", [])
                
                parts = []
                if bg:
                    parts.append(bg)
                for b in bullets:
                    parts.append(str(b))
                self._add_body_text(" ".join(parts))
            elif isinstance(summary_data, list):
                self._add_body_text(" ".join(str(item) for item in summary_data))
            else:
                self._add_body_text(str(summary_data))

        if "experience" in data and data["experience"]:
            self._add_heading("Professional Experience")
            self._add_experience_section(data["experience"])

        if "skills" in data and data["skills"]:
            self._add_heading("Skills")
            self._add_skills_section(data["skills"])

        if "certifications" in data and data["certifications"]:
            self._add_heading("Certifications")
            self._add_certifications_section(data["certifications"])

        if "education" in data and data["education"]:
            self._add_heading("Education")
            self._add_education_section(data["education"])

        if "projects" in data and data["projects"]:
            self._add_heading("PROJECTS")
            self._add_projects_section(data["projects"])

        return self.doc

    def _add_header_logo(self):
        """Add KPMG logo to the document header, right-aligned, using the valid PNG."""
        import os
        from docx.shared import Inches
        
        logo_path = os.path.join("assets", "kpmg_logo.png")
        if not os.path.exists(logo_path):
            return
            
        section = self.doc.sections[0]
        header = section.header
        
        # Make sure header has at least one paragraph
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = para.add_run()
        run.add_picture(logo_path, width=Inches(1.5))

    # ------------------------------------------------------------------
    # Section renderers
    # ------------------------------------------------------------------

    def _add_name_section(self, name: str, title: str = ""):
        """Add candidate name and title â€” matching the template formatting."""
        # Name
        para = self.doc.add_paragraph()
        run = para.add_run(name)
        self._set_run_font(run, size=self.NAME_SIZE, bold=True)
        self._set_paragraph_spacing(para, space_after=Pt(4), line_spacing=1.0)
        
        # Title
        if title:
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(title)
            self._set_run_font(title_run, size=self.TITLE_SIZE, italic=True)
            self._set_paragraph_spacing(title_para, space_after=Pt(12), line_spacing=1.0)

    def _add_heading(self, text: str):
        """Add a section heading with gray background shading."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        self._set_run_font(run, size=self.HEADING_SIZE, bold=True)
        self._set_paragraph_spacing(para, space_before=Pt(6), space_after=Pt(0),
                                     line_spacing=1.0)
        self._add_shading(para, self.SHADING_COLOR)

    def _add_body_text(self, text: str):
        """Add a regular body-text paragraph."""
        para = self.doc.add_paragraph()
        run = para.add_run(str(text))
        self._set_run_font(run, size=self.BODY_SIZE)
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._set_paragraph_spacing(para, space_after=Pt(0), line_spacing=1.0)

    def _add_experience_section(self, experience: list):
        """Add professional experience entries."""
        for job in experience:
            if isinstance(job, str):
                self._add_bullet_paragraph(job)
                continue

            # Job role + company + dates
            role = job.get("role", "")
            company = job.get("company", "")
            client = job.get("client", "")
            dates = job.get("dates", "")

            # Role line: "Role | Company/Client | Dates"
            role_line = role
            org_parts = []
            if company:
                org_parts.append(company)
            if client and client.lower() != company.lower():
                org_parts.append(client)
            if org_parts:
                if role_line:
                    role_line += f" | {' / '.join(org_parts)}"
                else:
                    role_line = " / ".join(org_parts)
            if dates:
                role_line += f" | {dates}"
            
            para = self.doc.add_paragraph()
            run = para.add_run(role_line)
            self._set_run_font(run, size=self.BODY_SIZE, bold=True)
            self._set_paragraph_spacing(para, space_before=Pt(4), space_after=Pt(0), line_spacing=1.0)

            # If there are no nested projects, output responsibilities directly under role line.
            nested_projects = job.get("projects", [])
            
            if not nested_projects:
                # "Role & Responsibilities" sub-heading
                rr = self.doc.add_paragraph()
                rr_run = rr.add_run("Role & Responsibilities")
                self._set_run_font(rr_run, size=self.BODY_SIZE, bold=True)
                self._set_paragraph_spacing(rr, space_after=Pt(0), line_spacing=1.0)

                # Responsibilities as bullet points
                responsibilities = job.get("responsibilities", [])
                if isinstance(responsibilities, str):
                    responsibilities = [responsibilities] if responsibilities else []
                for resp in responsibilities:
                    if resp:
                        self._add_bullet_paragraph(str(resp))

                # Technologies line
                technologies = job.get("technologies", [])
                if technologies:
                    if isinstance(technologies, list):
                        tech_str = ", ".join(str(t) for t in technologies if t)
                    else:
                        tech_str = str(technologies)
                    if tech_str:
                        tp = self.doc.add_paragraph()
                        # The template uses bold for "Technologies:" within paragraphs? 
                        # Wait, the template just output "Technologies: Azure...". The python extraction showed `[None, True, None]` which means bold applied selectively, or not.
                        run = tp.add_run(f"Technologies: {tech_str}")
                        self._set_run_font(run, size=self.BODY_SIZE)
                        self._set_paragraph_spacing(tp, space_after=Pt(4), line_spacing=1.0)

            # Nested projects under the job role
            nested_projects = job.get("projects", [])
            for p_idx, p in enumerate(nested_projects, 1):
                p_name = p.get("name", "")
                p_client = p.get("client", "")
                
                title_parts = []
                if p_name: title_parts.append(p_name)
                if p_client and p_client not in p_name: title_parts.append(p_client)
                p_title_str = " - ".join(title_parts)
                
                proj_head = self.doc.add_paragraph()
                p_run = proj_head.add_run(f"Project {p_idx}: {p_title_str}")
                self._set_run_font(p_run, size=self.BODY_SIZE, bold=True)
                self._set_paragraph_spacing(proj_head, space_before=Pt(6), space_after=Pt(0), line_spacing=1.0)
                
                # Check for project description and output it
                p_desc = p.get("description", "")
                if p_desc:
                    self._add_body_text(p_desc)
                
                prr = self.doc.add_paragraph()
                prr_run = prr.add_run("Role & Responsibilities")
                self._set_run_font(prr_run, size=self.BODY_SIZE, bold=True)
                self._set_paragraph_spacing(prr, space_after=Pt(0), line_spacing=1.0)

                p_resp = p.get("responsibilities", [])
                if isinstance(p_resp, str): p_resp = [p_resp]
                for r in p_resp:
                    if r: self._add_bullet_paragraph(str(r))
                    
                p_techs = p.get("technologies", [])
                if p_techs:
                    if isinstance(p_techs, list):
                        p_tech_str = ", ".join(str(t) for t in p_techs if t)
                    else:
                        p_tech_str = str(p_techs)
                    if p_tech_str:
                        ptp = self.doc.add_paragraph()
                        ptr = ptp.add_run(f"Technologies: {p_tech_str}")
                        self._set_run_font(ptr, size=self.BODY_SIZE, italic=True)
                        self._set_paragraph_spacing(ptp, space_before=Pt(0), space_after=Pt(4), line_spacing=1.0)


            # Blank separator between jobs
            self.doc.add_paragraph()

    def _add_skills_section(self, skills):
        """Add technical skills section."""
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                para = self.doc.add_paragraph()
                cat_run = para.add_run(f"{category}: ")
                self._set_run_font(cat_run, size=self.BODY_SIZE, bold=True)
                if isinstance(skill_list, list):
                    val_run = para.add_run(", ".join(str(s) for s in skill_list))
                else:
                    val_run = para.add_run(str(skill_list))
                self._set_run_font(val_run, size=self.BODY_SIZE)
                self._set_paragraph_spacing(para, space_after=Pt(2), line_spacing=1.0)

        elif isinstance(skills, list):
            for skill in skills:
                if isinstance(skill, dict):
                    cat = skill.get("category", "")
                    techs = skill.get("technologies", [])
                    para = self.doc.add_paragraph()
                    if cat:
                        cat_run = para.add_run(f"{cat}: ")
                        self._set_run_font(cat_run, size=self.BODY_SIZE, bold=True)
                    if isinstance(techs, list):
                        val_run = para.add_run(", ".join(str(t) for t in techs))
                    else:
                        val_run = para.add_run(str(techs))
                    self._set_run_font(val_run, size=self.BODY_SIZE)
                    self._set_paragraph_spacing(para, space_after=Pt(2), line_spacing=1.0)
                else:
                    self._add_bullet_paragraph(str(skill))

    def _add_certifications_section(self, certifications: list):
        """Add certifications section."""
        for cert in certifications:
            if isinstance(cert, dict):
                name = cert.get("name", "")
                cert_id = cert.get("id", "")
                date = cert.get("date", cert.get("earned", ""))
                expires = cert.get("expires", "")
                parts = [name]
                if cert_id:
                    parts.append(f"ID: {cert_id}")
                if date:
                    parts.append(f"Earned: {date}")
                if expires:
                    parts.append(f"Expires: {expires}")
                cert_text = " | ".join(parts)
            else:
                cert_text = str(cert)
            self._add_bullet_paragraph(cert_text)

    def _add_education_section(self, education: list):
        """Add education section."""
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                school = edu.get("school", edu.get("institution", ""))
                year = edu.get("year", edu.get("date", ""))
                edu_text = f"{degree} â€” {school}" if school else degree
                if year:
                    edu_text += f" ({year})"
            else:
                edu_text = str(edu)
            self._add_bullet_paragraph(edu_text)

    def _add_projects_section(self, projects: list):
        """Add projects section."""
        for project in projects:
            if isinstance(project, str):
                self._add_bullet_paragraph(project)
                continue

            proj_name = project.get("title", "") or project.get("name", "")
            company = project.get("company", "")
            client = project.get("client", "")
            dates = project.get("dates", "")

            # "Project#"
            para1 = self.doc.add_paragraph()
            run1 = para1.add_run("Project#")
            self._set_run_font(run1, size=self.BODY_SIZE, bold=True)
            self._set_paragraph_spacing(para1, space_after=Pt(0), line_spacing=1.0)
            
            # "Company Name - IBM"
            if company:
                para2 = self.doc.add_paragraph()
                run2 = para2.add_run(f"Company Name â€“ ")
                self._set_run_font(run2, size=self.BODY_SIZE, bold=True)
                run2_val = para2.add_run(company)
                self._set_run_font(run2_val, size=self.BODY_SIZE)
                self._set_paragraph_spacing(para2, space_after=Pt(0), line_spacing=1.0)

            # "Client Name: XYZ"
            if client:
                para3 = self.doc.add_paragraph()
                run3 = para3.add_run(f"Client Name: ")
                self._set_run_font(run3, size=self.BODY_SIZE, bold=True)
                run3_val = para3.add_run(client)
                self._set_run_font(run3_val, size=self.BODY_SIZE)
                self._set_paragraph_spacing(para3, space_after=Pt(0), line_spacing=1.0)

            # "Project Name: ProjTitle | Dates"
            title_line = proj_name
            if dates:
                title_line += f" | {dates}"
            
            para4 = self.doc.add_paragraph()
            run4 = para4.add_run(f"Project Name: ")
            self._set_run_font(run4, size=self.BODY_SIZE, bold=True)
            run4_val = para4.add_run(title_line)
            self._set_run_font(run4_val, size=self.BODY_SIZE)
            self._set_paragraph_spacing(para4, space_after=Pt(0), line_spacing=1.0)

            # Description
            description = project.get("description", "")
            if description:
                desc_label = self.doc.add_paragraph()
                lbl_run = desc_label.add_run("Description:")
                self._set_run_font(lbl_run, size=self.BODY_SIZE, bold=True)
                self._set_paragraph_spacing(desc_label, space_after=Pt(0), line_spacing=1.0)
                self._add_body_text(description)

            # Technologies
            technologies = project.get("technologies", [])
            if technologies:
                if isinstance(technologies, list):
                    tech_str = ", ".join(str(t) for t in technologies if t)
                else:
                    tech_str = str(technologies)
                if tech_str:
                    # Empty space before Technology
                    self.doc.add_paragraph()
                    tp = self.doc.add_paragraph()
                    run_tp = tp.add_run(f"Technology â€“ ")
                    self._set_run_font(run_tp, size=self.BODY_SIZE, bold=True)
                    run_tp_val = tp.add_run(tech_str)
                    self._set_run_font(run_tp_val, size=self.BODY_SIZE)
                    self._set_paragraph_spacing(tp, space_after=Pt(8), line_spacing=1.0)

            # Responsibilities
            responsibilities = project.get("responsibilities", [])
            if responsibilities:
                r_label = self.doc.add_paragraph()
                r_run = r_label.add_run("Responsibilities:")
                self._set_run_font(r_run, size=self.BODY_SIZE, bold=True)
                self._set_paragraph_spacing(r_label, space_after=Pt(0), line_spacing=1.0)
                
                if isinstance(responsibilities, str):
                    responsibilities = [responsibilities]
                for resp in responsibilities:
                    if resp:
                        self._add_bullet_paragraph(str(resp))
            
            # Blank separator between jobs
            self.doc.add_paragraph()

    # ------------------------------------------------------------------
    # Persistence
    # ------------------------------------------------------------------

    def save(self, output_path: str):
        """
        Save document to file.

        Args:
            output_path: Path to save DOCX file
        """
        try:
            self.doc.save(output_path)
        except Exception as e:
            raise Exception(f"Error saving document: {e}")
