import pdfplumber
from docx import Document
from typing import Tuple
import os


class ResumeParser:
    """Parser for extracting text from PDF and DOCX resume files"""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {e}")
        return text

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {e}")
        return text

    @staticmethod
    def parse(file_path: str) -> str:
        """
        Parse resume file (PDF or DOCX).
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == ".pdf":
            return ResumeParser.parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return ResumeParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")


class TemplateParser:
    """Parser for extracting template information from reference DOCX"""

    @staticmethod
    def extract_template_info(file_path: str) -> dict:
        """
        Extract template structure, styles, and formatting from reference DOCX.
        
        Args:
            file_path: Path to reference template DOCX
            
        Returns:
            Dictionary with template information
        """
        template_info = {
            "styles": {},
            "sections": [],
            "numbering": {},
            "tables": [],
            "logo": None,
            "margins": {},
            "fonts": {}
        }

        try:
            doc = Document(file_path)

            # Extract margins
            try:
                sections = doc.sections
                if sections:
                    section = sections[0]
                    template_info["margins"] = {
                        "top": section.top_margin.pt,
                        "bottom": section.bottom_margin.pt,
                        "left": section.left_margin.pt,
                        "right": section.right_margin.pt
                    }
            except:
                pass

            # Extract styles
            for style in doc.styles:
                try:
                    template_info["styles"][style.name] = {
                        "type": style.type,
                        "font_name": style.font.name if style.font.name else "Calibri",
                        "font_size": style.font.size.pt if style.font.size else 11,
                        "bold": style.font.bold,
                        "italic": style.font.italic,
                    }
                except:
                    pass

            # Extract paragraphs (section structure)
            for para in doc.paragraphs:
                # Truncate text massively to save LLM tokens (we only need structure/headers)
                t = para.text.strip()
                t = (t[:100] + "...") if len(t) > 100 else t
                if not t: continue

                section_info = {
                    "text_preview": t,
                    "level": para.style.name if para.style else "Normal",
                    "alignment": str(para.alignment),
                }
                
                # Get font details
                if para.runs:
                    run = para.runs[0]
                    section_info["font"] = {
                        "name": run.font.name or "Calibri",
                        "size": run.font.size.pt if run.font.size else 11,
                        "bold": run.font.bold,
                        "italic": run.font.italic,
                    }

                template_info["sections"].append(section_info)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_info = {
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "width": getattr(table, "width", None),
                }
                template_info["tables"].append(table_info)

            # Extract images (potential logo)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    template_info["logo"] = {
                        "exists": True,
                        "relation_id": rel.rId
                    }
                    break

        except Exception as e:
            raise Exception(f"Error extracting template: {e}")

        return template_info
